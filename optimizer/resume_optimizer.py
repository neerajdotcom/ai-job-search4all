import json
import logging
import os
import re
import shutil
import subprocess
import tempfile
from pathlib import Path

from google import genai
from docx import Document
from dotenv import load_dotenv

load_dotenv()

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
logger = logging.getLogger(__name__)

GEMINI_MODEL = "gemini-2.5-flash"


def build_system_prompt(profile) -> str:
    return (
        f"You are a resume optimizer for {profile.name}, a {profile.title} with "
        f"{profile.years_experience} years of {profile.industry_summary} experience. You optimize "
        "resumes for ATS and keyword match. You ONLY modify: Professional Summary, Core Competencies, "
        "and up to 3 bullet points per role. You never change dates, titles, companies, education, "
        "contact info, tools section, or any metrics. Never fabricate experience. Only reframe existing "
        "experience using JD language. Always respond with valid JSON only, no preamble, no markdown."
    )


OUTPUTS_DIR = Path(__file__).parent.parent / "outputs"

_METRIC_PATTERN = re.compile(r"\b\d+(?:\.\d+)?%|\$\d+(?:\.\d+)?[MBKmbk]?\b|\b\d+(?:\.\d+)?[xX]\b")


def extract_protected_metrics(text: str) -> set[str]:
    """Quantifiable figures (percentages, dollar amounts, multipliers) found in
    a resume's own text — derived per-resume rather than hardcoded, so the
    "never drop a metric" guard works for whatever numbers are actually in the
    candidate's own bullets."""
    return set(_METRIC_PATTERN.findall(text))


def extract_docx_text(docx_path: str) -> str:
    doc = Document(docx_path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


_client = None


def _get_client():
    global _client
    if _client is None:
        api_key = os.getenv("GEMINI_API_KEY", "").strip()
        if not api_key:
            raise RuntimeError("GEMINI_API_KEY not set")
        _client = genai.Client(api_key=api_key)
    return _client


import time as _time

# gemini-2.5-flash free tier is ~5 requests/min → space calls ~13s apart.
_MIN_INTERVAL = 13.0
_last_call = [0.0]


def _throttle():
    elapsed = _time.time() - _last_call[0]
    if elapsed < _MIN_INTERVAL:
        _time.sleep(_MIN_INTERVAL - elapsed)
    _last_call[0] = _time.time()


# Once Gemini reports the free-tier *daily* request cap is hit, every later
# call this run fails the same way — retrying with 20-40s backoff per job just
# burns time for nothing. This flag short-circuits the rest of the run as soon
# as that's detected, instead of doomed-retrying through every remaining job.
_daily_quota_exhausted = [False]


def _is_daily_quota_error(msg: str) -> bool:
    return "PerDay" in msg or "free_tier_requests" in msg


def _generate_with_retry(client, prompt, max_retries=3):
    """Call Gemini, retrying transient 429/503 with backoff."""
    for attempt in range(max_retries + 1):
        _throttle()
        try:
            response = client.models.generate_content(model=GEMINI_MODEL, contents=prompt)
            return response.text
        except Exception as exc:
            msg = str(exc)
            if "RESOURCE_EXHAUSTED" in msg and _is_daily_quota_error(msg):
                _daily_quota_exhausted[0] = True
                logger.error("Gemini daily free-tier quota exhausted — skipping Gemini for the rest of this run")
                raise
            transient = ("503" in msg or "UNAVAILABLE" in msg
                         or "429" in msg or "RESOURCE_EXHAUSTED" in msg)
            if transient and attempt < max_retries:
                wait = 20 * (attempt + 1)
                logger.warning("Gemini transient error (attempt %d) — retrying in %ds",
                               attempt + 1, wait)
                _time.sleep(wait)
                continue
            raise


def call_llm(resume_text: str, jd: str, profile, company: str = "", title: str = "") -> dict:
    client = _get_client()

    archetype_lines = "\n".join(f"    * {k}: {v}" for k, v in profile.archetypes.items())
    role_keys = list(profile.roles.keys())
    role_keys_str = ", ".join(role_keys) if role_keys else "role1, role2, role3"
    protected_metrics = extract_protected_metrics(resume_text)
    protected_metrics_str = ", ".join(sorted(protected_metrics)) if protected_metrics else "none found"

    prompt = (
        f"{build_system_prompt(profile)}\n\n"
        f"{profile.context}\n"
        f"Target role: {title} at {company}\n\n"
        f"Here is the resume:\n\n{resume_text}\n\n"
        f"Here is the job description:\n\n{jd}\n\n"
        "First classify this JD into the single best-fit framing archetype, then "
        "bias the optimized summary, competencies, and bullets toward that "
        "archetype's emphasis (without fabricating anything). Archetypes:\n"
        f"{archetype_lines}\n\n"
        "Return ONLY a JSON object with these keys:\n"
        "- match_score (integer 0-100)\n"
        "- missing_keywords (list of strings)\n"
        "- strong_matches (list of strings)\n"
        f"- chosen_archetype (string: exactly one of {', '.join(profile.archetypes)})\n"
        "- optimized_summary (string)\n"
        "- optimized_competencies (list of strings)\n"
        f"- optimized_bullets (object with keys: {role_keys_str} — each a list of max 3 strings)\n"
        "- cover_note_variants (list of 2-3 objects, each with 'angle' (one of: "
        "domain-fit, delivery-track-record, growth-story) and 'text' (a concise 4-6 "
        "sentence cover note written from that angle, ready to paste into an "
        f"application, signed '{profile.name}'). Each variant should genuinely "
        "lead with a different hook, not just reword the same note.)\n"
        "- screening_answers (list of objects, each with 'question' and 'answer', covering the "
        "common application questions: notice period, why this company/role, total years of "
        "experience, willingness to relocate, current location, and expected CTC. "
        "Base every answer ONLY on the candidate facts above — never fabricate.)\n"
        f"Preserve all metrics exactly as written ({protected_metrics_str}). No preamble. No markdown."
    )

    raw = (_generate_with_retry(client, prompt) or "").strip()
    # Strip markdown code fences if the model wraps despite instructions
    raw = re.sub(r"^```(?:json)?\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)
    result = json.loads(raw)

    # Backward-compat: downstream (digest, dashboard, snapshot) still reads a
    # single cover_note. Keep it populated from the first variant so nothing
    # breaks if a consumer hasn't been updated, while the full variants list is
    # also carried through.
    variants = result.get("cover_note_variants")
    if isinstance(variants, list) and variants and isinstance(variants[0], dict):
        result.setdefault("cover_note", variants[0].get("text", ""))
    return result


def _missing_metrics(text: str, protected_metrics: set[str]) -> list[str]:
    return [metric for metric in protected_metrics if metric not in text]


# Headless LibreOffice is the renderer — it's the only cross-platform way to
# turn the patched DOCX into a PDF that faithfully preserves the resume's
# template formatting (no Python lib reproduces Word layout without rebuilding
# it from scratch). It's preinstalled on GitHub's ubuntu-latest runners and on
# most dev machines; if it's missing we fall back to the DOCX, never crash.
_SOFFICE_TIMEOUT = 120


def _render_pdf(docx_path: Path) -> Path | None:
    """Render a DOCX to PDF via headless LibreOffice. Returns the PDF path on
    success, or None if LibreOffice isn't available or the conversion fails —
    callers fall back to the DOCX so a missing renderer never breaks the run."""
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        logger.warning("LibreOffice (soffice) not found — cannot render PDF, keeping DOCX")
        return None

    pdf_path = docx_path.with_suffix(".pdf")
    # A private, throwaway user-profile dir avoids clashing with any LibreOffice
    # the user already has open (a shared profile would make headless conversion
    # hang or fail), and keeps sequential calls from stepping on each other.
    profile_dir = Path(tempfile.mkdtemp(prefix="lo_profile_"))
    try:
        proc = subprocess.run(
            [
                soffice,
                f"-env:UserInstallation=file://{profile_dir}",
                "--headless", "--convert-to", "pdf",
                "--outdir", str(docx_path.parent),
                str(docx_path),
            ],
            check=True, capture_output=True, timeout=_SOFFICE_TIMEOUT, text=True,
        )
    except (subprocess.CalledProcessError, subprocess.TimeoutExpired, OSError) as exc:
        logger.warning("LibreOffice PDF conversion failed for %s: %s", docx_path.name, exc)
        return None
    finally:
        shutil.rmtree(profile_dir, ignore_errors=True)

    # soffice notoriously exits 0 even when it couldn't load the source, so the
    # file-existence check below is the real success gate — surface its captured
    # output to make a silent failure diagnosable on the runner.
    if not pdf_path.exists():
        logger.warning(
            "LibreOffice exited cleanly but produced no PDF for %s — stdout=%r stderr=%r",
            docx_path.name, (proc.stdout or "").strip(), (proc.stderr or "").strip(),
        )
        return None
    return pdf_path


def _find_paragraph_index(doc: Document, search_text: str):
    """Locate a heading paragraph by text. Prefers heading-shaped matches
    (exact text equality, or a short paragraph whose stripped text starts
    with the query) over a substring hit anywhere in the doc, because a
    naive substring search matches the query inside a paragraph body — a
    real bug caught in production: `search_text="Experience"` matched the
    Professional Summary body ("...Cross-Functional Program Coordination.
    Experience...") on one candidate's CV and returned the summary
    paragraph as the "Experience heading," which then poisoned the section-
    heading size reference and misclassified every downstream bullet."""
    query = search_text.strip().lower()

    # Pass 1: exact match (case-insensitive), ignoring surrounding whitespace.
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().lower() == query:
            return i
    # Pass 2: heading-shaped starts-with match on a short paragraph.
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if len(text) <= 60 and text.lower().startswith(query):
            return i
    # Pass 3: legacy substring fallback (still needed for older templates).
    for i, p in enumerate(doc.paragraphs):
        if query in p.text.lower():
            return i
    return None


def _replace_paragraph_text(paragraph, new_text: str):
    """Replace all runs in a paragraph with a single run containing new_text."""
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)


def patch_docx(
    source_path: str,
    result: dict,
    profile,
    company_name: str,
    title: str = "",
) -> tuple[Path, Path, list[str]]:
    OUTPUTS_DIR.mkdir(exist_ok=True)
    safe_company = re.sub(r"[^\w\-]", "", company_name.replace(" ", ""))
    safe_title = re.sub(r"[^\w\-]", "", title.replace(" ", ""))[:40]
    slug = f"{safe_company}-{safe_title}" if safe_title else safe_company

    # The editable DOCX is the working file we patch; the PDF rendered from it
    # is the resume that actually gets submitted/attached. The DOCX is kept as
    # an editable companion the user can tweak by hand before applying.
    docx_path = OUTPUTS_DIR / f"{profile.resume_slug}_Resume-{slug}.docx"

    # Work from an independent copy — never touch the original
    shutil.copy2(source_path, docx_path)

    quality_warnings: list[str] = []

    protected_metrics = extract_protected_metrics(extract_docx_text(source_path))

    doc = Document(str(docx_path))

    # --- Summary ---
    summary_idx = _find_paragraph_index(doc, "Professional Summary") or _find_paragraph_index(doc, "Summary")
    if summary_idx is not None:
        # The actual summary text is the next non-empty paragraph after the heading
        for i in range(summary_idx + 1, min(summary_idx + 4, len(doc.paragraphs))):
            if doc.paragraphs[i].text.strip():
                _replace_paragraph_text(doc.paragraphs[i], result["optimized_summary"])
                break
    else:
        logger.warning("Could not locate Professional Summary section in %s — left unpatched", docx_path.name)
        quality_warnings.append("Professional Summary section not found — left unpatched")

    # --- Core Competencies ---
    # Try common heading variants in specificity order — most CVs use one of
    # "Core Competencies" / "Competencies" / "Skills Summary" / "Skills" /
    # "Key Skills" for the skills-and-tools section. First match wins.
    comp_idx = (
        _find_paragraph_index(doc, "Core Competencies")
        or _find_paragraph_index(doc, "Competencies")
        or _find_paragraph_index(doc, "Skills Summary")
        or _find_paragraph_index(doc, "Key Skills")
        or _find_paragraph_index(doc, "Skills")
    )
    if comp_idx is not None:
        for i in range(comp_idx + 1, min(comp_idx + 4, len(doc.paragraphs))):
            if doc.paragraphs[i].text.strip():
                _replace_paragraph_text(
                    doc.paragraphs[i],
                    " | ".join(result["optimized_competencies"]),
                )
                break
    else:
        logger.warning("Could not locate Core Competencies section in %s — left unpatched", docx_path.name)
        quality_warnings.append("Core Competencies section not found — left unpatched")

    # --- Role bullets ---
    # role_key -> match keywords comes from the candidate's own profile
    # (profile.roles), not a hardcoded employer list, so this patches
    # whatever companies are in THIS resume.
    role_map = profile.roles
    all_role_keywords = [kw for kws in role_map.values() for kw in kws]

    # Bullet detection needs a real signal: many resume templates (including
    # the reference templates this tool ships with) render bullets via Word's
    # numbering XML (w:numPr), so python-docx's paragraph.text never contains
    # a literal "•"/"-" character, and the bullet paragraphs' style name is
    # just "Normal" — the same as everything else. The reliable signal is
    # relative font size: a top-level SECTION heading ("PROFESSIONAL
    # EXPERIENCE", "EDUCATION", ...) is bold and set in a *larger* font than
    # the bold sub-headings used *within* a role's bullet block (e.g. "Sprint
    # Execution & Workload Management"), which are themselves larger than
    # nothing — both are bold, but only the section heading matches the
    # reference size below. Bullets are simply the non-bold paragraphs.
    section_idx = _find_paragraph_index(doc, "Professional Experience") or _find_paragraph_index(doc, "Experience")
    section_heading_size = None
    section_heading_text = None
    if section_idx is not None and doc.paragraphs[section_idx].runs:
        section_heading_size = doc.paragraphs[section_idx].runs[0].font.size
        section_heading_text = doc.paragraphs[section_idx].text.strip()

    def _is_section_heading(p) -> bool:
        if not p.runs or not p.runs[0].bold:
            return False
        text = p.text.strip()
        # First: recognize the reference "PROFESSIONAL EXPERIENCE" heading
        # itself, plus other conventional section markers, purely by text.
        # These are the strings that unambiguously start a new section on
        # any resume template — treating them as section headings by name
        # works even when font-size info is missing (a real production bug:
        # `size is None or sh_size is None` used to fire True on every bold
        # paragraph, misclassifying role sub-headings as section boundaries
        # and losing all bullets in that role).
        _WELL_KNOWN_SECTIONS = {
            "professional experience", "experience", "work experience",
            "education", "academics", "certifications", "publications",
            "awards", "additional information", "skills", "skills summary",
            "core competencies", "competencies", "key skills", "tools",
            "professional summary", "summary",
        }
        if text.lower() in _WELL_KNOWN_SECTIONS:
            return True
        size = p.runs[0].font.size
        # Only trip the size compare when we HAVE both sizes to compare.
        # A missing size (paragraph inherits from the paragraph-mark default)
        # is not enough evidence on its own to call something a section
        # heading — that decision now belongs to the well-known-text check
        # above, not to a "no data" fallback that was over-aggressive.
        if size is not None and section_heading_size is not None:
            return size >= section_heading_size
        return False

    def _is_role_subheading(p) -> bool:
        return bool(p.runs) and bool(p.runs[0].bold) and not _is_section_heading(p)

    for role_key, keywords in role_map.items():
        bullets = result.get("optimized_bullets", {}).get(role_key, [])
        if not bullets:
            continue

        # Find the heading paragraph for this company. Prefer a bold match
        # first — role headings are bold on every template we've seen, and
        # this filter prevents a real bug: after the Professional Summary
        # is patched (above), the tailored summary text often mentions the
        # candidate's current employer by name, so a naive substring hit on
        # keywords like "arrise" would land on the Summary body (a
        # non-bold paragraph) instead of the actual role heading further
        # down. Fall back to any match only if no bold match exists.
        role_idx = None
        for i, p in enumerate(doc.paragraphs):
            if p.runs and p.runs[0].bold and any(kw in p.text.lower() for kw in keywords):
                role_idx = i
                break
        if role_idx is None:
            for i, p in enumerate(doc.paragraphs):
                if any(kw in p.text.lower() for kw in keywords):
                    role_idx = i
                    break

        if role_idx is None:
            logger.warning("Could not locate role section for '%s' in DOCX — skipping bullets", role_key)
            quality_warnings.append(f"Role section '{role_key}' not found — bullets left unpatched")
            continue

        # Collect EVERY bullet paragraph under the role (until the next role
        # or section). We used to stop collecting at 3 — but the DOCX still
        # physically contained bullets 4+, so the patched output showed 3
        # tailored bullets followed by whatever originals were in place,
        # producing duplicated/stale content. We now collect all of them,
        # overwrite the first N with the tailored bullets, and CLEAR the
        # remainder in place so the reader sees only the tailored set.
        bullet_paragraphs = []
        for i in range(role_idx + 1, len(doc.paragraphs)):
            p = doc.paragraphs[i]
            text = p.text.strip()
            if not text:
                continue
            if any(kw in text.lower() for kw in all_role_keywords) or _is_section_heading(p):
                break
            if _is_role_subheading(p):
                continue
            bullet_paragraphs.append(p)

        if not bullet_paragraphs:
            logger.warning("Found role section for '%s' but no bullet paragraphs under it — left unpatched", role_key)
            quality_warnings.append(f"Role section '{role_key}' found but no bullets detected — left unpatched")
            continue

        # Overwrite the first len(bullets) paragraphs; blank the rest so the
        # role section shows only the tailored set. Blanking preserves the
        # paragraph's numbering XML / bullet-list membership so the surviving
        # bullets keep their template-driven bullet character — safer than
        # deleting paragraphs, which python-docx does not support cleanly.
        n = min(len(bullet_paragraphs), len(bullets))
        for bp, new_bullet in zip(bullet_paragraphs[:n], bullets[:n]):
            original = bp.text
            # Guard: never drop a protected metric
            for metric in protected_metrics:
                if metric in original and metric not in new_bullet:
                    new_bullet = new_bullet.rstrip(".") + f", achieving {metric} improvement."
            _replace_paragraph_text(bp, new_bullet)
        # Blank any excess original bullets left over after the tailored set,
        # so the tailored resume doesn't ship with stale bullets underneath.
        # EXCEPT: if a surplus bullet contains a protected metric (a %/$/Nx
        # figure the guard above insists must survive verbatim), keep it —
        # blanking would drop the metric, which is the exact behavior the
        # "never drop a protected metric" guard exists to prevent.
        for bp in bullet_paragraphs[n:]:
            if any(m in bp.text for m in protected_metrics):
                continue
            _replace_paragraph_text(bp, "")

    doc.save(str(docx_path))

    missing = _missing_metrics(extract_docx_text(str(docx_path)), protected_metrics)
    if missing:
        logger.warning("Protected metric(s) missing from optimized resume: %s", ", ".join(missing))
        quality_warnings.append(f"Protected metric(s) missing from output: {', '.join(missing)}")

    # Render the patched DOCX to PDF — the submittable deliverable. If
    # LibreOffice is missing or conversion fails, fall back to the DOCX as the
    # submit artifact so the run never breaks over a renderer hiccup.
    pdf_path = _render_pdf(docx_path)
    if pdf_path is None:
        quality_warnings.append("PDF render failed — DOCX used as the submit file (install LibreOffice)")
        ats_path = docx_path
    else:
        ats_path = pdf_path

    logger.info("Saved submit version → %s", ats_path)
    logger.info("Saved editable DOCX  → %s", docx_path)
    return ats_path, docx_path, quality_warnings


def optimize_resume(docx_path: str, jd: str, profile, company_name: str, title: str = "") -> dict:
    """
    Main entry point. Always called for jobs that already cleared the canonical
    Groq match-score threshold in main.py — so this function always writes the
    DOCX. (It used to also gate on Gemini's own internal match_score >= 60, but
    Gemini and Groq score the same JD independently and can disagree — e.g. Groq
    62% vs Gemini 55% for the same posting — which silently dropped the resume
    for a job already shown as "qualifying" in the digest. Groq's score is the
    one that decided this job belongs here, so it's the only gate that matters.)

    Returns the full LLM JSON result, including a tailored cover_note and
    screening_answers for a fast-apply pack.
    """
    if _daily_quota_exhausted[0]:
        raise RuntimeError("Gemini daily quota exhausted earlier this run — optimization skipped to save time.")

    logger.info("Starting resume optimization for company: %s", company_name)

    resume_text = extract_docx_text(docx_path)
    logger.info("Extracted %d characters from resume", len(resume_text))

    result = call_llm(resume_text, jd, profile, company_name, title)
    score = result.get("match_score", 0)
    logger.info("Gemini's own match score for %s: %d%% (informational only)", company_name, score)

    ats_path, review_path, quality_warnings = patch_docx(docx_path, result, profile, company_name, title)
    result["ats_output"] = str(ats_path)
    result["review_output"] = str(review_path)
    result["quality_warnings"] = quality_warnings

    return result


if __name__ == "__main__":
    import sys
    from candidate_profile.loader import load_profile, EXAMPLE_PROFILE_PATH

    profile = load_profile(sys.argv[2] if len(sys.argv) > 2 else str(EXAMPLE_PROFILE_PATH))
    resume_file = sys.argv[1] if len(sys.argv) > 1 else profile.resume_path
    sample_jd = (
        f"Looking for a {profile.title} with agile delivery, sprint planning, "
        "stakeholder management, and Jira experience."
    )
    sample_company = "TestCo"

    output = optimize_resume(resume_file, sample_jd, profile, sample_company)

    print("\n--- Optimization Result ---")
    print(json.dumps(output, indent=2))
